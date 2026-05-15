from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"D:\School Work & Docs\SEM_3.2 Units\Project\SecureVote_Project_Report_Draft0.docx")

PROJECT_TITLE = "A Secure and Verifiable Student E-Voting System"
COURSE = "CCS/CSE 2328 Project I"
DATE = "May 2026"

ACCENT = "0F766E"
NAVY = "0F172A"
BLUE = "1D4ED8"
PALE = "F8FAFC"
GRAY = "64748B"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcw = tc_pr.first_child_found_in("w:tcW")
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tc_pr.append(tcw)
    tcw.set(qn("w:w"), str(int(width_inches * 1440)))
    tcw.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color="FFFFFF")
        shade_cell(cell, header_fill)
        if widths:
            set_cell_width(cell, widths[i])
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i, item in enumerate(row):
            set_cell_text(cells[i], item)
            if widths:
                set_cell_width(cells[i], widths[i])
            if r % 2 == 0:
                shade_cell(cells[i], PALE)
    doc.add_paragraph()
    return table


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_callout(doc, title, body, fill="ECFDF5", border="10B981"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), border)
        borders.append(elem)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(border)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(9)
    doc.add_paragraph()


def configure_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = styles[name]
        st.font.name = "Aptos Display"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        st.font.color.rgb = RGBColor.from_string(NAVY)
        st.font.bold = True
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(11)
    for style_name in ["List Bullet", "List Number"]:
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.size = Pt(10)
    try:
        code = styles.add_style("CodeSmall", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        code = styles["CodeSmall"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(8)
    code.font.color.rgb = RGBColor.from_string("334155")
    footer = sec.footer.paragraphs[0]
    footer.text = f"{PROJECT_TITLE} | Draft 0 Report | {DATE}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(GRAY)


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("SECUREVOTE")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(PROJECT_TITLE)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Eliminating queues, empowering every student, and securing every vote")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_paragraph()
    add_callout(
        doc,
        "Report Type",
        "Draft 0 full project documentation report including Chapters 1-7, references, and appendices.",
        fill="F0FDFA",
        border=ACCENT,
    )
    add_table(
        doc,
        ["Item", "Details"],
        [
            ("Course", COURSE),
            ("Project Area", "Secure full-stack e-voting system for student elections"),
            ("Technology Stack", "Flask REST API, HTML/CSS/Vanilla JavaScript, SQLite"),
            (
                "Security Focus",
                "OTP, bcrypt, AES-GCM, SHA-256 vote receipts, append-only votes, hash-chained audit logs",
            ),
            ("Date", DATE),
        ],
        widths=[1.7, 5.8],
        header_fill=ACCENT,
    )
    doc.add_page_break()


def front_matter(doc):
    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "This project presents a secure and verifiable student e-voting system designed to replace slow, queue-based, and delegate-centred campus election processes with a direct digital voting platform. The system allows eligible students to register, authenticate using password and one-time password verification, vote for candidates by position, receive verification hashes, and review selected election results. Administrators can create elections, register candidates before polls open, open and close elections, view results, export reports, inspect voters, and monitor audit logs.",
    )
    add_para(
        doc,
        "The implemented prototype follows a three-tier architecture consisting of a responsive HTML/CSS/JavaScript frontend, a Python Flask REST API backend, and a SQLite database. Security features include bcrypt password hashing, AES-GCM encrypted vote storage, SHA-256 vote receipts, logical and database-level vote immutability, role-based access control, duplicate vote prevention per position, and hash-chained audit logs for tamper detection.",
    )
    add_heading(doc, "Table of Contents", 1)
    add_numbered(
        doc,
        [
            "Chapter 1: Introduction",
            "Chapter 2: Literature Review",
            "Chapter 3: Methodology",
            "Chapter 4: System Analysis",
            "Chapter 5: System Design",
            "Chapter 6: Implementation and Testing",
            "Chapter 7: Results and Conclusion",
            "References",
            "Appendices",
        ],
    )
    add_heading(doc, "List of Abbreviations", 1)
    add_table(
        doc,
        ["Abbreviation", "Meaning"],
        [
            ("AES", "Advanced Encryption Standard"),
            ("API", "Application Programming Interface"),
            ("DBMS", "Database Management System"),
            ("EVM", "Electronic Voting Machine"),
            ("MFA", "Multi-Factor Authentication"),
            ("OTP", "One-Time Password"),
            ("RBAC", "Role-Based Access Control"),
            ("REST", "Representational State Transfer"),
            ("SHA", "Secure Hash Algorithm"),
            ("SQL", "Structured Query Language"),
            ("UAT", "User Acceptance Testing"),
        ],
        widths=[1.5, 5.5],
        header_fill=BLUE,
    )
    doc.add_page_break()


def chapter1(doc):
    add_heading(doc, "Chapter 1: Introduction", 1)
    add_heading(doc, "1.1 Background of the Study", 2)
    add_para(
        doc,
        "Student elections are a central part of campus governance because they allow learners to choose representatives who speak for their academic, social, welfare, and financial interests. In many institutions, however, the election process still depends on physical queues, paper ballots, or delegate-based arrangements where representatives vote on behalf of a wider student body. These approaches can reduce direct participation, slow down result publication, and create doubts about whether the final outcome reflects the actual opinion of students.",
    )
    add_para(
        doc,
        "The SecureVote project responds to this challenge by providing a web-based voting system where students can participate directly from any supported device during an active election window. The system is not only concerned with convenience; it also addresses security and trust through authentication, encrypted vote storage, vote verification receipts, duplicate vote prevention, and tamper-evident audit logs.",
    )
    add_heading(doc, "1.2 Problem Statement", 2)
    add_para(
        doc,
        "The current voting approach is affected by long queues, limited participation, paper-heavy counting, delayed results, and weak independent verification. Where a delegate model is used, many students do not personally cast a vote, which creates a democratic gap. Manual counting also increases the risk of error, disputes, and delayed announcement of results.",
    )
    add_callout(
        doc,
        "Core Problem",
        "There is a need for a fast, accessible, secure, and verifiable platform that allows every eligible student to vote directly while maintaining election integrity.",
        fill="EFF6FF",
        border=BLUE,
    )
    add_heading(doc, "1.3 Aim of the Project", 2)
    add_para(
        doc,
        "The aim of this project is to design and implement a secure and verifiable student e-voting system that allows students to vote independently, quickly, and safely in campus elections.",
    )
    add_heading(doc, "1.4 Research Objectives", 2)
    add_numbered(
        doc,
        [
            "Analyse the shortcomings of the current student voting approach.",
            "Design an e-voting system that improves security, transparency, and accessibility.",
            "Implement a working full-stack prototype using Flask, HTML/CSS/JavaScript, and SQLite.",
            "Provide vote verification, audit logging, and duplicate vote prevention mechanisms.",
            "Recommend future integration strategies for campus election deployment.",
        ],
    )
    add_heading(doc, "1.5 Scope of the Project", 2)
    add_para(
        doc,
        "The system focuses on student council and faculty-level elections where registered students vote for candidates in predefined positions such as President, Finance, and Academics. It supports registration, login, OTP verification, election setup, candidate setup before polls open, voting, verification, results viewing, audit review, and report export. National elections, multi-campus federation elections, and legally binding public elections are outside the prototype scope.",
    )
    add_heading(doc, "1.6 Significance of the Study", 2)
    add_bullets(
        doc,
        [
            "Students gain a direct, accessible voting channel without physical queues.",
            "Election administrators gain faster tallying, cleaner records, and better audit visibility.",
            "The institution gains a prototype that demonstrates how digital election integrity can be improved using practical security controls.",
            "The project contributes an applied software artefact that can be evaluated, improved, and adapted for campus governance.",
        ],
    )
    doc.add_page_break()


def chapter2(doc):
    add_heading(doc, "Chapter 2: Literature Review", 1)
    add_heading(doc, "2.1 Introduction", 2)
    add_para(
        doc,
        "Electronic voting research examines how digital systems can improve participation, speed, accuracy, and transparency while still protecting ballot secrecy and election integrity. This chapter reviews paper-based voting, electronic voting machines, and web-based voting systems, then identifies the gap that motivates a campus-specific secure e-voting prototype.",
    )
    add_heading(doc, "2.2 Related Systems", 2)
    add_table(
        doc,
        ["System", "Strengths", "Limitations for this Project"],
        [
            (
                "Helios Voting",
                "Open-audit web voting, cryptographic ballot receipts, end-to-end verifiability.",
                "Stronger cryptographic model than needed for a simple campus prototype; requires careful voter education.",
            ),
            (
                "Polyas Online Voting",
                "Commercial online voting, cloud deployment, multi-factor authentication, device support.",
                "Commercial and less suited for a student-built prototype where internal learning and customization are required.",
            ),
            (
                "iVote / Government Internet Voting",
                "Designed for remote accessibility and large election administration.",
                "Government-scale systems are complex, costly, and have documented security concerns.",
            ),
            (
                "Manual campus voting",
                "Simple to understand and easy to observe physically.",
                "Queues, delayed counting, lower participation, and limited independent verification.",
            ),
        ],
        widths=[1.5, 3.0, 3.0],
        header_fill=BLUE,
    )
    add_heading(doc, "2.3 Weaknesses of Existing Approaches", 2)
    add_table(
        doc,
        ["Area", "Observed Weakness", "SecureVote Response"],
        [
            ("Security", "Paper ballots can be misplaced or manipulated; online systems can be attacked if poorly designed.", "Password hashing, OTP verification, encrypted vote storage, protected APIs, and audit logs."),
            ("Transparency", "Manual counting can be opaque; closed systems require trust in administrators.", "Voters receive SHA-256 verification hashes and admins see tamper-evident audit status."),
            ("Accessibility", "Physical polling locations and queues reduce participation.", "Students can vote through a responsive browser interface during the active election window."),
            ("Vote verification", "Many systems do not allow a voter to confirm that a vote was recorded.", "A student-only vote verification page checks whether a vote hash exists."),
            ("Campus suitability", "General-purpose systems do not always match small campus election workflows.", "The prototype supports fixed student positions, admin setup lock, and role-specific dashboards."),
        ],
        widths=[1.3, 3.0, 3.2],
        header_fill=NAVY,
    )
    add_heading(doc, "2.4 Research Gap", 2)
    add_para(
        doc,
        "The reviewed systems show that secure electronic voting is possible, but many existing solutions are either too manual, too expensive, too general, or too complex for a campus prototype. A gap exists for a lightweight student election system that combines direct participation, authentication, encrypted vote storage, receipt-based verification, structured voting by position, and tamper-evident administration.",
    )
    add_heading(doc, "2.5 Proposed Solution from the Review", 2)
    add_para(
        doc,
        "The proposed solution is a secure student e-voting system using standard web technologies and simplified but correct security mechanisms. Instead of attempting to replicate a national election-grade cryptographic platform, the project demonstrates practical controls that support confidentiality, integrity, availability, authentication, and accountability at prototype level.",
    )
    doc.add_page_break()


def chapter3(doc):
    add_heading(doc, "Chapter 3: Methodology", 1)
    add_heading(doc, "3.1 Research Approach", 2)
    add_para(
        doc,
        "The project uses an applied design-science approach. The study constructs a working software artefact and evaluates whether the artefact addresses the identified election problems. This approach fits the project because the main contribution is a functional system supported by security analysis, user workflow design, and testing results.",
    )
    add_heading(doc, "3.2 Development Methodology", 2)
    add_para(
        doc,
        "Agile development was selected because the project required iterative improvement. Requirements such as structured voting by position, role-based navigation, live results, setup locking, and closed-election result selection were refined through repeated testing and feedback.",
    )
    add_table(
        doc,
        ["Agile Stage", "Activities Completed"],
        [
            ("Planning", "Defined users, roles, election workflow, security expectations, and prototype scope."),
            ("Design", "Prepared database schema, API structure, dashboards, voting flow, and security services."),
            ("Development", "Implemented Flask routes, services, SQLite tables, frontend pages, and JavaScript API calls."),
            ("Testing", "Tested authentication, OTP, voting, duplicate prevention, audit tampering, and role restrictions."),
            ("Refinement", "Improved UI/UX, responsive layout, role-specific navigation, result filtering, and admin controls."),
        ],
        widths=[1.5, 5.8],
        header_fill=ACCENT,
    )
    add_heading(doc, "3.3 Data Collection Methods", 2)
    add_para(
        doc,
        "The presentation proposed questionnaires, user acceptance testing sessions, and review of secondary literature. For the prototype, these tools support requirement validation and help confirm that the system addresses real student and administrator needs.",
    )
    add_bullets(
        doc,
        [
            "Questionnaires: gather student views on current elections, usability expectations, trust, and willingness to use online voting.",
            "Interviews: collect administrator requirements on election setup, candidate management, result publication, and audit needs.",
            "Observation: document problems in physical or delegate-based election workflows, including queues and counting delays.",
            "UAT sessions: allow selected users to test registration, login, OTP, voting, verification, and result viewing.",
        ],
    )
    add_heading(doc, "3.4 Ethical Considerations", 2)
    add_para(
        doc,
        "Participants should be informed about the purpose of the study, the voluntary nature of participation, and how responses are used. Test data should avoid unnecessary personal information. In a production deployment, personal data would require institutional approval, privacy controls, and secure retention policies.",
    )
    add_heading(doc, "3.5 Development Tools", 2)
    add_table(
        doc,
        ["Layer", "Tool/Technology", "Purpose"],
        [
            ("Frontend", "HTML, CSS, Vanilla JavaScript", "Responsive pages, dashboards, forms, polling, alerts, and API calls."),
            ("Backend", "Python Flask", "REST API routes, sessions, authentication, voting, admin management, and results."),
            ("Database", "SQLite", "Prototype storage for users, elections, candidates, votes, OTP sessions, audit logs, and Merkle roots."),
            ("Security", "bcrypt, AES-GCM, SHA-256, hash chains", "Password protection, encrypted votes, receipt hashes, and tamper evidence."),
            ("Testing", "Browser, terminal, SQLite inspection", "Functional, UI, and security validation."),
        ],
        widths=[1.2, 2.0, 4.0],
        header_fill=BLUE,
    )
    doc.add_page_break()


def chapter4(doc):
    add_heading(doc, "Chapter 4: System Analysis", 1)
    add_heading(doc, "4.1 Existing System Analysis", 2)
    add_para(
        doc,
        "The existing voting process relies on manual or delegate-centred activities. Students either queue physically or depend on representatives to vote on their behalf. Election officials manually verify voters, issue ballots, count votes, and announce results. This workflow is familiar but slow and vulnerable to participation gaps, counting disputes, and weak individual verification.",
    )
    add_heading(doc, "4.2 Proposed System Analysis", 2)
    add_para(
        doc,
        "The proposed system introduces a secure online workflow. Students register, log in, verify OTP, view an active election, vote once per position, receive verification hashes, and can later verify recorded vote hashes. Admins manage elections and candidates before polls open, open or close elections, view selected results, export reports, and inspect audit logs.",
    )
    add_heading(doc, "4.3 Functional Requirements", 2)
    add_table(
        doc,
        ["ID", "Requirement", "Implementation Status"],
        [
            ("FR1", "Student registration and login with OTP verification.", "Implemented."),
            ("FR2", "Admin election management: create, open, close, and list elections.", "Implemented with setup locked while an election is open."),
            ("FR3", "Candidate management by position.", "Implemented with fixed positions: President, Finance, Academics."),
            ("FR4", "Secure vote casting.", "Implemented with AES-GCM encrypted storage and SHA-256 receipt generation."),
            ("FR5", "Duplicate vote prevention.", "Implemented per user, election, and position."),
            ("FR6", "Vote verification.", "Implemented for authenticated students only."),
            ("FR7", "Results and report export.", "Implemented with active/closed election result selection."),
            ("FR8", "Audit log review and tamper detection.", "Implemented using hash-chained audit logs and integrity checks."),
        ],
        widths=[0.7, 4.0, 2.8],
        header_fill=NAVY,
    )
    add_heading(doc, "4.4 Non-Functional Requirements", 2)
    add_table(
        doc,
        ["Quality Attribute", "Requirement", "Prototype Response"],
        [
            ("Security", "Protect credentials and votes.", "bcrypt passwords, OTP, AES-GCM votes, protected routes."),
            ("Integrity", "Prevent vote changes and detect tampering.", "Append-only vote logic, SQLite triggers, hash-chained logs."),
            ("Usability", "Simple workflow and responsive interface.", "Modern dashboards, role-based navigation, mobile-aware pages."),
            ("Availability", "Support election activity during voting period.", "Local prototype runs on Flask; production hosting is recommended."),
            ("Verifiability", "Allow voters and admins to confirm integrity.", "Vote hash verification and audit chain checks."),
            ("Maintainability", "Keep code modular.", "Routes, services, models, utilities, frontend JS modules, and schema separated."),
        ],
        widths=[1.4, 2.8, 3.1],
        header_fill=ACCENT,
    )
    add_heading(doc, "4.5 Main Actors", 2)
    add_bullets(
        doc,
        [
            "Student: registers, logs in, verifies OTP, votes, views profile, checks vote hash, and views selected election results.",
            "Admin: creates elections, adds candidates before opening, opens/closes elections, views results, exports reports, manages voter roles, and reviews audit logs.",
            "System: encrypts votes, generates hashes, checks duplicates, stores audit records, calculates statistics, and validates integrity.",
        ],
    )
    doc.add_page_break()


def chapter5(doc):
    add_heading(doc, "Chapter 5: System Design", 1)
    add_heading(doc, "5.1 Proposed System Overview", 2)
    add_para(
        doc,
        "SecureVote is a browser-based student e-voting platform that combines user authentication, structured ballots, encrypted vote storage, receipt-based verification, and role-specific administration. The implemented system follows a clean three-tier architecture: presentation layer, application layer, and data layer.",
    )
    add_heading(doc, "5.2 Three-Tier Architecture", 2)
    add_table(
        doc,
        ["Tier", "Components", "Responsibility"],
        [
            ("Presentation Layer", "HTML pages, CSS styling, Vanilla JavaScript modules", "Displays login, OTP, dashboard, vote, profile, verify, admin, results, audit, and voter management interfaces."),
            ("Application Layer", "Flask app, auth routes, voting routes, admin routes, services, decorators", "Validates requests, enforces roles, manages sessions, processes votes, computes results, and records audit logs."),
            ("Data Layer", "SQLite database, schema, triggers, seed data", "Stores users, elections, candidates, votes, OTP sessions, audit logs, and Merkle roots."),
        ],
        widths=[1.5, 2.7, 3.3],
        header_fill=BLUE,
    )
    add_heading(doc, "5.3 Database Design", 2)
    add_para(
        doc,
        "The database schema supports the required election entities and security records. Votes include encrypted vote data, encrypted key metadata, vote hashes, voter blind hashes, and position information. A unique constraint prevents a student from voting twice for the same position in the same election.",
    )
    add_table(
        doc,
        ["Table", "Main Fields", "Purpose"],
        [
            ("users", "id, name, email, password_hash, role, reg_number, course, has_voted", "Stores student/admin accounts and profile details."),
            ("elections", "id, title, start_date, end_date, status, is_active, created_by", "Stores election periods and active/closed state."),
            ("candidates", "id, name, election_id, position", "Stores candidate names by election and position."),
            ("votes", "id, user_id, candidate_id, election_id, position, encrypted_vote, encrypted_key, vote_hash", "Stores encrypted append-only vote records and receipts."),
            ("audit_logs", "id, action, user_id, timestamp, previous_hash, current_hash, merkle_root", "Stores hash-chained security and activity logs."),
            ("otp_sessions", "id, user_id, otp_code, expiry_time, is_used", "Stores temporary OTP records for login verification."),
            ("merkle_roots", "id, election_id, merkle_root, computed_at", "Stores final vote root for closed-election integrity checks."),
        ],
        widths=[1.3, 3.5, 2.7],
        header_fill=NAVY,
    )
    add_heading(doc, "5.4 Voting Workflow", 2)
    add_numbered(
        doc,
        [
            "Student logs in using email and password.",
            "System generates OTP and verifies the submitted OTP before creating a full voting session.",
            "Student opens the voting page and views candidates grouped by position.",
            "Student selects one candidate per position and submits all selected votes.",
            "Backend validates eligibility, active election status, role, candidate-position match, and duplicate vote constraints.",
            "Each vote is encrypted, hashed, inserted into the votes table, and recorded in the audit chain.",
            "Student receives one or more SHA-256 verification hashes as vote receipts.",
            "Results are calculated for the active election or for a selected closed election.",
        ],
    )
    add_heading(doc, "5.5 Security Design", 2)
    add_table(
        doc,
        ["Security Control", "Design Purpose", "Implementation"],
        [
            ("Password hashing", "Protect passwords if database is exposed.", "bcrypt hashes stored in users table."),
            ("OTP verification", "Add a second login factor.", "OTP sessions with expiry and used status."),
            ("Role-based access control", "Separate student and admin functions.", "Flask decorators and frontend role-aware navigation."),
            ("AES-GCM vote encryption", "Keep stored votes unreadable in plaintext.", "Vote payload encrypted before database insertion."),
            ("SHA-256 vote hash", "Give voters a receipt for verification.", "Vote hash generated and returned after successful voting."),
            ("Append-only votes", "Protect vote records after casting.", "Application logic plus SQLite triggers prevent UPDATE/DELETE."),
            ("Audit hash chain", "Detect modified audit rows.", "Each log stores previous_hash and current_hash."),
            ("Merkle root on close", "Summarize final vote hashes for election integrity.", "Computed and stored when an election closes."),
        ],
        widths=[1.8, 2.8, 2.8],
        header_fill=ACCENT,
    )
    add_heading(doc, "5.6 User Interface Design", 2)
    add_para(
        doc,
        "The frontend was upgraded into a modern dashboard-style interface using a dark glassmorphism visual system with teal/green accents. Navigation is role-aware: unauthenticated users see Login and Register; students see Dashboard, Vote, Profile, Verify Vote, and Logout; admins see Dashboard, Results, Audit Logs, Voters, and Logout.",
    )
    doc.add_page_break()


def chapter6(doc):
    add_heading(doc, "Chapter 6: Implementation and Testing", 1)
    add_heading(doc, "6.1 Development Environment", 2)
    add_para(
        doc,
        "The backend was implemented in Python using the Flask framework and REST API structure. The frontend was implemented using HTML, CSS, and Vanilla JavaScript. SQLite was selected as the database engine because it is lightweight, easy to seed, and suitable for a working academic prototype.",
    )
    add_table(
        doc,
        ["Area", "Technology", "Reason"],
        [
            ("Backend", "Python Flask", "Simple REST routing, sessions, decorators, and service structure."),
            ("Frontend", "HTML, CSS, Vanilla JavaScript", "No framework dependency; compatible with Flask static serving."),
            ("Database", "SQLite", "Lightweight local database suitable for prototype setup and testing."),
            ("Security libraries", "bcrypt, PyCryptodome/cryptographic utilities, hashlib", "Password hashing, AES encryption, RSA wrapping, and SHA-256 hashing."),
            ("Development tools", "VS Code, browser, terminal, SQLite tools", "Coding, testing, database inspection, and UI verification."),
        ],
        widths=[1.4, 2.4, 3.5],
        header_fill=BLUE,
    )
    add_heading(doc, "6.2 Implemented Modules", 2)
    add_bullets(
        doc,
        [
            "Authentication module: registration, login, OTP generation, OTP verification, session management, and logout.",
            "Voting module: active election retrieval, structured candidate grouping, vote submission, duplicate prevention, receipt generation, and verification.",
            "Admin module: election creation, candidate addition, election opening/closing, results, audit logs, voter roles, and export report.",
            "Security services: password hashing, vote encryption, vote hashing, audit hash chaining, Merkle root generation, and integrity checks.",
            "Frontend modules: reusable navigation, role-aware pages, loading states, alerts, dashboard polling, result selection, and responsive UI.",
        ],
    )
    add_heading(doc, "6.3 Test Data", 2)
    add_table(
        doc,
        ["Category", "Sample Data"],
        [
            ("User accounts", "Admin account and student accounts with emails, passwords, registration numbers, and courses."),
            ("Election data", "Student Council election with active and closed states."),
            ("Candidates", "Candidates grouped under President, Finance, and Academics."),
            ("Security tests", "Duplicate vote attempts, SQL update attempts on votes, modified audit log rows, and role-crossing access attempts."),
        ],
        widths=[1.7, 5.7],
        header_fill=NAVY,
    )
    add_heading(doc, "6.4 Test Cases and Results", 2)
    add_table(
        doc,
        ["ID", "Feature", "Action", "Expected Result", "Result"],
        [
            ("TC-01", "Authentication", "Login with valid credentials and OTP.", "Session created and user redirected by role.", "Passed"),
            ("TC-02", "Vote casting", "Student selects candidates per position.", "Encrypted votes stored and hashes returned.", "Passed"),
            ("TC-03", "Duplicate prevention", "Same student votes again for same position.", "Request blocked and logged.", "Passed"),
            ("TC-04", "Vote encryption", "Inspect vote record in database.", "Plain candidate choice is not stored as readable text.", "Passed"),
            ("TC-05", "Vote immutability", "Attempt SQL UPDATE/DELETE on votes.", "Trigger rejects modification.", "Passed"),
            ("TC-06", "Vote verification", "Submit valid vote hash.", "System confirms hash exists.", "Passed"),
            ("TC-07", "Tamper detection", "Modify audit log action manually.", "Audit chain reports tampering.", "Passed"),
            ("TC-08", "RBAC", "Admin attempts student-only verify route.", "Access denied.", "Passed"),
        ],
        widths=[0.65, 1.35, 2.3, 2.4, 0.8],
        header_fill=ACCENT,
    )
    add_heading(doc, "6.5 Testing Summary", 2)
    add_para(
        doc,
        "All eight core functional and security tests passed during prototype validation. Testing confirmed that the system supports login and OTP, records encrypted votes, prevents duplicate votes per position, blocks direct vote modification, verifies vote hashes, detects audit log tampering, and restricts access by role. Remaining production concerns include real email/SMS OTP delivery, HTTPS hosting, managed key storage, database backups, and formal penetration testing.",
    )
    doc.add_page_break()


def chapter7(doc):
    add_heading(doc, "Chapter 7: Results and Conclusion", 1)
    add_heading(doc, "7.1 Achievements and Lessons Learnt", 2)
    add_para(
        doc,
        "The project achieved its main objective by delivering a working secure student e-voting prototype aligned with the research aim and objectives. The system moved beyond a static design and became a functional platform with authentication, OTP verification, role-based access, election management, structured voting, encrypted vote storage, vote verification, audit logging, and modern dashboards.",
    )
    add_table(
        doc,
        ["Objective", "Achievement"],
        [
            ("Analyse shortcomings of current voting", "The project identified queues, delegate voting, delayed counting, weak verification, and limited transparency as major issues."),
            ("Design a secure and transparent system", "A three-tier architecture, secure workflow, database schema, UI pages, APIs, and security controls were designed."),
            ("Implement a working prototype", "A Flask, SQLite, HTML/CSS/JavaScript system was built and tested end-to-end."),
            ("Support verification and integrity", "Vote hashes, AES-GCM encryption, append-only votes, hash-chained audit logs, and Merkle roots were implemented."),
            ("Improve participation and usability", "The responsive UI supports student registration, direct voting by position, dashboards, profile, verification, and result selection."),
        ],
        widths=[2.4, 5.1],
        header_fill=BLUE,
    )
    add_para(
        doc,
        "Important lessons were learnt during development. First, election security must be designed into the workflow from the beginning; it cannot be added only at the interface level. Second, hiding a button in the UI is not sufficient security, so backend validation was also required for actions such as adding candidates after an election is open. Third, usability affects trust: students and administrators understand the system better when results, audit status, and voting state are clearly displayed. Finally, prototype cryptography must be described honestly: the system demonstrates correct security principles, while full production deployment requires stronger operational controls such as managed keys, HTTPS, real OTP delivery, and formal audits.",
    )
    add_heading(doc, "7.2 Conclusions", 2)
    add_para(
        doc,
        "The SecureVote prototype demonstrates that a campus election process can be made faster, more accessible, and more transparent using a full-stack web application with practical security controls. Students can vote directly without queues, and administrators can manage elections, candidates, results, voters, and audit logs from a controlled dashboard. The system supports one vote per position per election, which better matches real student council elections than a single-vote model.",
    )
    add_para(
        doc,
        "The project also shows that vote integrity can be strengthened through layered controls. Password hashing protects credentials, OTP improves login assurance, AES-GCM protects stored vote data, SHA-256 receipts support voter verification, database triggers protect vote immutability, and hash-chained audit logs expose tampering attempts. Although the prototype is not a final production election platform, it is a solid proof of concept that meets the academic aim of building and demonstrating a secure, verifiable student e-voting system.",
    )
    add_heading(doc, "7.3 Recommendations for Future Work", 2)
    add_bullets(
        doc,
        [
            "Integrate real email or SMS OTP delivery instead of simulated console OTP output.",
            "Deploy over HTTPS with secure cookies, CSRF protection, hardened CORS, and institution-managed hosting.",
            "Move encryption keys and secrets into a managed key vault or hardware security module.",
            "Add stronger end-to-end verifiable election protocols such as homomorphic tallying or mixnets for higher-stakes deployments.",
            "Introduce formal candidate approval workflows, election scheduling, nomination windows, and automatic open/close timing.",
            "Add PDF report generation with official election summaries, audit status, turnout, and candidate results.",
            "Conduct wider usability testing with students, administrators, and election committee members.",
            "Implement backup, recovery, monitoring, and incident response procedures for production use.",
            "Add accessibility testing to improve keyboard navigation, screen reader support, and low-bandwidth usability.",
        ],
    )
    doc.add_page_break()


def references(doc):
    add_heading(doc, "References", 1)
    refs = [
        "Adida, B. (2008). Helios: Web-based open-audit voting. Proceedings of the 17th USENIX Security Symposium, 335-348.",
        "Alvarez, R. M., & Hall, T. E. (2008). Electronic elections: The perils and promises of digital democracy. Princeton University Press.",
        "Beck, K., Beedle, M., van Bennekum, A., Cockburn, A., Cunningham, W., Fowler, M., & Thomas, D. (2001). Manifesto for agile software development. Agile Alliance. https://agilemanifesto.org",
        "Chaum, D., Carback, R., Clark, J., Essex, A., Popoveniuc, S., Rivest, R. L., Ryan, P. Y. A., Shen, E., & Sherman, A. T. (2008). Scantegrity II: End-to-end verifiability for optical scan election systems using invisible ink confirmation codes. IEEE Security & Privacy, 6(3), 14-21.",
        "Creswell, J. W. (2014). Research design: Qualitative, quantitative, and mixed methods approaches (4th ed.). SAGE Publications.",
        "Gritzalis, D. A. (2002). Principles and requirements for a secure e-voting system. Computers & Security, 21(6), 539-556.",
        "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information systems research. MIS Quarterly, 28(1), 75-105.",
        "Kitchenham, B. (2004). Procedures for performing systematic reviews. Keele University Technical Report TR/SE-0401.",
        "Mercuri, R. T. (2002). A better ballot box? IEEE Spectrum, 39(10), 46-50.",
        "National Institute of Standards and Technology. (2001). Advanced Encryption Standard (AES) (FIPS Publication No. 197). https://doi.org/10.6028/NIST.FIPS.197",
        "National Institute of Standards and Technology. (2015). Secure Hash Standard (SHS) (FIPS Publication No. 180-4). https://doi.org/10.6028/NIST.FIPS.180-4",
        "Pallets Projects. (2024). Flask documentation. https://flask.palletsprojects.com/",
        "Provos, N., & Mazieres, D. (1999). A future-adaptable password scheme. Proceedings of the USENIX Annual Technical Conference, 81-91.",
        "SQLite Consortium. (2024). SQLite documentation. https://www.sqlite.org/docs.html",
        "Springall, D., Finkenauer, T., Durumeric, Z., Kitcat, J., Hursti, H., MacAlpine, M., & Halderman, J. A. (2014). Security analysis of the Estonian internet voting system. Proceedings of the ACM SIGSAC Conference on Computer and Communications Security, 703-715.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_page_break()


def appendices(doc):
    add_heading(doc, "Appendices", 1)
    add_heading(doc, "Appendix A: User Manual", 2)
    add_heading(doc, "A.1 Starting the System", 3)
    add_numbered(
        doc,
        [
            "Open a terminal in the project folder.",
            "Create and activate a Python virtual environment.",
            "Install dependencies using pip install -r requirements.txt.",
            "Seed the database using python -m backend.seed if a fresh demo database is required.",
            "Start the server using python -m backend.app.",
            "Open http://127.0.0.1:5000 in a browser.",
        ],
    )
    add_heading(doc, "A.2 Student Guide", 3)
    add_numbered(
        doc,
        [
            "Register using name, email, registration number, course, and password.",
            "Log in with email and password.",
            "Enter the OTP generated by the system. In the prototype, OTP is simulated for local testing.",
            "Open the voting page during an active election.",
            "Select one candidate under each available position.",
            "Submit the vote and save the verification hash shown on the confirmation page.",
            "Use Verify Vote to check whether the hash exists in the system.",
            "Use the dashboard dropdown to view closed election results when no election is active.",
        ],
    )
    add_heading(doc, "A.3 Admin Guide", 3)
    add_numbered(
        doc,
        [
            "Log in using an admin account and complete OTP verification.",
            "Create an election while no election is open.",
            "Add candidates under President, Finance, and Academics before opening the election.",
            "Open the election when setup is complete. Candidate and election setup becomes locked while voting is in progress.",
            "Close the election after voting ends. The system computes integrity data and records audit activity.",
            "Use Results to select an active or closed election and view counts, percentages, turnout, and leading candidates.",
            "Use Audit Logs to check the hash chain and tamper status.",
            "Use Voters to review registered users and adjust roles where permitted.",
            "Use Export Report to download a report for the selected election.",
        ],
    )
    add_heading(doc, "A.4 Troubleshooting", 3)
    add_table(
        doc,
        ["Problem", "Likely Cause", "Action"],
        [
            ("Cannot log in", "Wrong credentials or OTP not verified.", "Check email/password, request a new OTP, and retry before expiry."),
            ("No vote page content", "No active election exists.", "Admin must open an election before students can vote."),
            ("Add Candidate hidden", "Election is already open.", "Close the active election before changing setup."),
            ("Audit status shows tampered", "Audit log row was modified.", "Restore from backup or reseed test data for a clean demo."),
            ("Results show zero", "No active election or no closed election selected.", "Select a closed election from the dropdown or open an election."),
        ],
        widths=[1.8, 2.6, 3.0],
        header_fill=NAVY,
    )
    doc.add_page_break()

    add_heading(doc, "Appendix B: Data Collection Tools", 2)
    add_heading(doc, "B.1 Student Questionnaire", 3)
    add_para(
        doc,
        "Purpose: To collect student views on current election challenges, online voting acceptance, usability expectations, and security concerns. Suggested response scale: Strongly Agree, Agree, Neutral, Disagree, Strongly Disagree.",
    )
    add_table(
        doc,
        ["No.", "Question"],
        [
            ("1", "The current student election process is convenient for most students."),
            ("2", "Long queues discourage students from voting."),
            ("3", "I trust the current method of counting and announcing results."),
            ("4", "I would prefer to vote directly from my phone or computer."),
            ("5", "I would trust an online system more if it used OTP verification."),
            ("6", "Receiving a vote verification hash would increase my confidence in the system."),
            ("7", "The voting interface should group candidates by position."),
            ("8", "I would use an online voting system if it protected my privacy."),
        ],
        widths=[0.6, 6.7],
        header_fill=BLUE,
    )
    add_heading(doc, "B.2 Admin Interview Guide", 3)
    add_bullets(
        doc,
        [
            "How are student elections currently created, supervised, and closed?",
            "What records are needed before, during, and after an election?",
            "What are the common causes of disputes in student elections?",
            "Who should be allowed to create elections, add candidates, and publish results?",
            "What audit information should be visible to election administrators?",
            "What reports are required after an election is completed?",
        ],
    )
    add_heading(doc, "B.3 Observation Checklist", 3)
    add_table(
        doc,
        ["Area Observed", "What to Record"],
        [
            ("Queue management", "Average waiting time, crowding, and accessibility issues."),
            ("Voter verification", "How eligibility is checked and whether duplicates can occur."),
            ("Ballot handling", "How ballots are issued, stored, counted, and protected."),
            ("Result announcement", "Time taken to publish results and how disputes are handled."),
            ("Transparency", "Whether students can independently verify their vote or the final count."),
        ],
        widths=[2.0, 5.4],
        header_fill=ACCENT,
    )
    doc.add_page_break()

    add_heading(doc, "Appendix C: Project Schedule", 2)
    add_table(
        doc,
        ["Phase", "Duration", "Activities", "Deliverable"],
        [
            ("Initiation", "Week 1", "Topic selection, problem identification, project aim and objectives.", "Approved project concept."),
            ("Literature Review", "Weeks 2-3", "Review e-voting systems, security requirements, and related work.", "Chapter 2 draft."),
            ("Analysis", "Week 4", "Collect requirements and define functional/non-functional requirements.", "System analysis chapter."),
            ("Design", "Week 5", "Design architecture, database, workflow, APIs, and UI wireframes.", "System design chapter."),
            ("Implementation Sprint 1", "Weeks 6-7", "Build authentication, OTP, database schema, and base frontend.", "Working login/register flow."),
            ("Implementation Sprint 2", "Weeks 8-9", "Build election management, structured voting, encryption, audit logs, and results.", "Functional secure voting prototype."),
            ("Testing and Refinement", "Week 10", "Run functional, UI, security, and UAT tests; improve UI/UX.", "Test report and final prototype."),
            ("Documentation", "Week 11", "Prepare full report, Chapter 7, references, and appendices.", "Draft 0 documentation report."),
        ],
        widths=[1.5, 1.0, 3.2, 1.8],
        header_fill=NAVY,
    )

    add_heading(doc, "Appendix D: Project Budget", 2)
    add_table(
        doc,
        ["Item", "Description", "Estimated Cost (KES)"],
        [
            ("Internet/data", "Research, development, testing, and deployment preparation.", "3,000"),
            ("Hosting allowance", "Prototype hosting or campus server preparation.", "3,500"),
            ("Domain/SSL allowance", "Domain and certificate planning for production deployment.", "2,000"),
            ("Email/SMS OTP testing", "Allowance for future real OTP gateway testing.", "2,500"),
            ("Printing and binding", "Draft and final report printing.", "2,500"),
            ("User testing support", "Refreshments/logistics for UAT participants.", "3,000"),
            ("Contingency", "Unexpected costs during testing or presentation.", "3,000"),
            ("Total", "Estimated project budget.", "19,500"),
        ],
        widths=[1.8, 3.8, 1.8],
        header_fill=ACCENT,
    )
    add_heading(doc, "Appendix E: API Endpoint Summary", 2)
    add_table(
        doc,
        ["Area", "Endpoint", "Purpose"],
        [
            ("Auth", "POST /register", "Create student account."),
            ("Auth", "POST /login", "Validate password and generate OTP."),
            ("Auth", "POST /verify-otp", "Verify OTP and create session."),
            ("Voting", "GET /election", "Retrieve active election and candidates."),
            ("Voting", "POST /vote", "Submit votes by position."),
            ("Voting", "GET /verify-vote?hash=", "Verify student vote receipt hash."),
            ("Admin", "POST /admin/create-election", "Create a closed election setup."),
            ("Admin", "POST /admin/add-candidate", "Add candidate before election opens."),
            ("Admin", "POST /admin/open-election", "Open election for voting."),
            ("Admin", "POST /admin/close-election", "Close election and record integrity data."),
            ("Admin", "GET /admin/results", "Return counts and percentages."),
            ("Admin", "GET /admin/audit-logs", "Return audit log entries and integrity status."),
            ("Admin", "GET /admin/export-report", "Export selected election report."),
        ],
        widths=[1.1, 2.4, 3.8],
        header_fill=BLUE,
    )


def main():
    doc = Document()
    configure_document(doc)
    cover(doc)
    front_matter(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    chapter7(doc)
    references(doc)
    appendices(doc)
    doc.core_properties.title = PROJECT_TITLE
    doc.core_properties.subject = "Draft 0 full project documentation report"
    doc.core_properties.author = "SecureVote Project Team"
    doc.core_properties.keywords = "e-voting, Flask, SQLite, AES, SHA-256, audit logs, student elections"
    doc.save(OUT)
    print(f"Created {OUT}")


if __name__ == "__main__":
    main()
